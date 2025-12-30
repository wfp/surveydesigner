import datetime
import io

import pandas as pd
from django.conf import settings
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR
from docx.shared import Cm, Pt


class DocConversion:
    def __init__(self, xlsx_file, languages=None):
        self.xlsx = xlsx_file
        self.doc = Document()
        self.languages = languages or []
        self.module_num = 0
        self.submodule_num = 0
        self.question_num = 0
        self.form_number = 0
        self.current_module = None
        self.current_submodule = None
        self.available_languages = settings.LANGUAGES
        self.available_languages_lookup = [lang[0] for lang in settings.LANGUAGES]

    def _remove_trailing_whitespace(self):
        for paragraph in self.doc.paragraphs:
            if len(paragraph.text) == 0:
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None

    def _add_select(self, table, question, language):
        choices = question["choices"]
        choices_sheet = pd.read_excel(self.xlsx, sheet_name="choices")
        code, label = language

        # Filter for the specific choice list
        choice_list = choices_sheet.loc[choices_sheet["list_name"] == choices].copy()

        label_column = f"label::{label} ({code})"
        choice_list[label_column] = choice_list[label_column].fillna("-")

        # Construct the choice_label series
        if choices not in ("villages", "gps", "regions"):
            choice_label = (
                choice_list["name"].astype(str)
                + "  "
                + choice_list[label_column].astype(str)
            )
        else:
            choice_label = choice_list[label_column]

        # Iterate preserving order - use .values to get numpy array or .tolist()
        for choice in choice_label.tolist():
            choice = " ".join(choice.split())
            paragraph = table.cell(0, 1).add_paragraph()
            run = paragraph.add_run()
            run.text = "\u2610 " + choice
            font = run.font
            font.size = Pt(9)

    def _add_relevant_and_constraint(self, table, relevant, constraint):
        paragraph = table.cell(0, 1).add_paragraph()

        def add_text(message, bold=False):
            run = paragraph.add_run(message)
            run.font.size = Pt(9)
            run.italic = True
            run.bold = bold

        add_text("(")

        if relevant != "nan":
            add_text("skip logic: ")
            add_text(relevant, bold=True)
            if constraint != "nan":
                add_text(", ")

        if constraint != "nan":
            add_text("validation: ")
            add_text(constraint, bold=True)

        add_text(")")

    def _add_question(self, question, language):
        type_ = question["type"].split(" ", 1)[0]
        label = question["label"]
        hint = question["hint"]
        name = question["name"]
        relevant = str(question["relevant"])
        constraint = str(question["constraint"])

        # Increment question number only for actual questions
        if type_ not in (
            "note",
            "begin_group",
            "begin_repeat",
            "end_repeat",
            "start",
            "end",
            "today",
            "deviceid",
        ):
            self.question_num += 1

        table = self.doc.add_table(1, 2)
        paragraph = table.cell(0, 0).add_paragraph()

        # Always use the current module, submodule, and question numbers
        paragraph.add_run(
            f"{self.module_num}.{self.submodule_num}.{self.question_num} "
        ).bold = True
        paragraph.add_run(f"{name}, {type_}: ").italic = True
        paragraph.add_run(label).bold = True

        if str(hint) not in ("nan", "-"):
            run = paragraph.add_run(" Hint: " + hint)
            run.italic = True
            # Arabic in italic doesn't render correctly in some editors
            if "Arabic" in language:
                run.italic = False

        if relevant != "nan" or constraint != "nan":
            self._add_relevant_and_constraint(table, relevant, constraint)

        if type_ in ("select_one", "select_multiple"):
            self._add_select(table, question, language)

    def _add_submodule_heading(self, paragraph, row):
        label = row["label"]
        relevant = row["relevant"]

        # Increment submodule number and reset question number
        self.submodule_num += 1
        self.question_num = 0

        self.doc.add_heading(f"{self.module_num}.{self.submodule_num} {label}", 2)

        def add_relevant_text(message, bold=False):
            run = paragraph.add_run(message)
            run.font.size = Pt(9)
            run.font.color.theme_color = MSO_THEME_COLOR.ACCENT_5
            run.italic = True
            run.bold = bold

        if (relevant_str := str(relevant)) != "nan":
            paragraph = self.doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            add_relevant_text("(skip logic: ")
            add_relevant_text(relevant_str, bold=True)
            add_relevant_text(")")

    def _add_module_heading(self, row):
        label = row["label"]

        # Increment module number and reset submodule and question numbers
        self.module_num += 1
        self.submodule_num = 0
        self.question_num = 0

        self.doc.add_heading(f"{self.module_num} {label}", 1)

    def _process_group(self, paragraph, row):
        name = row["name"]
        if name.endswith("_module"):
            self._add_module_heading(row)
        elif name.endswith("_submodule"):
            self._add_submodule_heading(paragraph, row)

    def _add_repeat_section_end(self):
        table = self.doc.add_table(1, 2)
        paragraph = table.cell(0, 0).add_paragraph()
        paragraph.add_run("End repeat section.").italic = True

    def _add_repeat_section_start(self, row):
        table = self.doc.add_table(1, 2)
        paragraph = table.cell(0, 0).add_paragraph()
        repeat_section_label = row["label"]
        paragraph.add_run("Begin repeat section: ").italic = True
        paragraph.add_run(repeat_section_label).bold = True

    def _add_note_section(self, paragraph, row):
        paragraph.add_run(row["label"]).bold = True
        if str(row["hint"]) not in ("nan", "-"):
            paragraph.add_run(" Hint: " + row["hint"]).italic = True

    def _process_row(self, row, language):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)

        type_ = row["type"]

        if type_ in ("start", "end", "today", "deviceid"):
            return
        elif type_ == "begin_repeat":
            self._add_repeat_section_start(row)
        elif type_ == "end_repeat":
            self._add_repeat_section_end()
        elif type_ == "begin_group":
            self._process_group(p, row)
        elif type_ == "note":
            self._add_note_section(p, row)
        else:
            self._add_question(row, language)

    def _get_xls_rows(self):
        survey_sheet = pd.read_excel(self.xlsx, sheet_name="survey")
        return survey_sheet.loc[
            ~survey_sheet["type"].isin(
                ["end_group", "username", "simserial", "calculate"]
            )
        ]

    def _add_title(self, language):
        code, label = language
        settings_sheet = pd.read_excel(self.xlsx, sheet_name="settings")
        title = settings_sheet["form_title"][0]

        if self.form_number > 0:
            self.doc.add_page_break()

        self.doc.add_heading(f"{title}: {label} ({code})", 0)
        self.doc.add_paragraph().add_run(
            "This file was generated on " + datetime.datetime.now().strftime("%d-%b-%Y")
        ).italic = True

    def _format_margins(self, margin_cm):
        for section in self.doc.sections:
            section.top_margin = Cm(margin_cm)
            section.bottom_margin = Cm(margin_cm)
            section.left_margin = Cm(margin_cm)
            section.right_margin = Cm(margin_cm)

    def _process_form(self, xls_rows, language):
        # Explicitly reset all counters at the start of each form
        self.module_num = 0
        self.submodule_num = 0
        self.question_num = 0

        self._add_title(language)

        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(9)

        code, label = language

        def handle_nan(cell):
            return cell_str if (cell_str := str(cell)) != "nan" else "-"

        for i in xls_rows.index:
            row_type = xls_rows["type"][i].split(" ", 1)
            row = {
                "type": row_type[0],
                "choices": row_type[1] if len(row_type) == 2 else None,
                "label": handle_nan(xls_rows[f"label::{label} ({code})"][i]),
                "hint": handle_nan(xls_rows[f"hint::{label} ({code})"][i]),
                "name": xls_rows["name"][i],
                "relevant": xls_rows["relevant"][i],
                "constraint": xls_rows["constraint"][i],
            }
            self._process_row(row, language)

    def process(self):
        self._format_margins(1)
        xls_rows = self._get_xls_rows()
        for code in self.languages:
            try:
                index = self.available_languages_lookup.index(code)
            except ValueError:
                continue
            else:
                language = self.available_languages[index]
                self._process_form(xls_rows, language)
                self.form_number += 1

        self._remove_trailing_whitespace()

    def run(self):
        self.process()
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer
